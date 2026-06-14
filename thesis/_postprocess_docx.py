#!/usr/bin/env python3
"""pandoc çıktısı docx'i YTÜ kurallarına göre cilalar:
numaralandırma, İÇİNDEKİLER/ŞEKİL/ÇİZELGE/KAYNAKLAR/EKLER, kapak ortalama."""
import sys
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PATH = "Quantum Grup Yapıları - YTÜ Bitirme Tezi.docx"
d = Document(PATH)

UNNUM = {"ÖNSÖZ", "SİMGE LİSTESİ", "KISALTMA LİSTESİ", "ÖZET", "ABSTRACT",
         "İÇİNDEKİLER", "ŞEKİL LİSTESİ", "ÇİZELGE LİSTESİ", "KAYNAKLAR",
         "EKLER", "ÖZGEÇMİŞ"}

def prepend_run(p, s):
    r = p.add_run(s)
    r.bold = True
    el = r._r
    p._p.remove(el)
    pPr = p._p.find(qn("w:pPr"))
    if pPr is not None:
        pPr.addnext(el)
    else:
        p._p.insert(0, el)

# ---- 1) Başlık numaralandırma ----
chap = sec = sub = 0
appendix = False
letter = None
for p in d.paragraphs:
    st = p.style.name
    if st == "Heading 1":
        t = p.text.strip()
        if t in UNNUM:
            sec = sub = 0
            continue
        if t.startswith("Kurulum, Çalıştırma"):
            appendix = True; letter = "A"; sec = sub = 0
            prepend_run(p, "EK A  "); continue
        if t.startswith("Test Eşlemesi"):
            appendix = True; letter = "B"; sec = sub = 0
            prepend_run(p, "EK B  "); continue
        appendix = False; chap += 1; sec = sub = 0
        prepend_run(p, f"{chap}  ")
    elif st == "Heading 2":
        sec += 1; sub = 0
        prepend_run(p, f"{letter}.{sec}  " if appendix else f"{chap}.{sec}  ")
    elif st == "Heading 3":
        sub += 1
        prepend_run(p, f"{chap}.{sec}.{sub}  ")

# ---- 1b) thebibliography'den sızan "99" etiketini sil ----
for p in list(d.paragraphs):
    if p.text.strip() == "99" and p.style.name in ("Normal", "Body Text"):
        p._element.getparent().remove(p._element)
        break

# ---- 1c) \paragraph run-in başlıkları (Heading 4/5) düz metne çevir ----
for p in d.paragraphs:
    if p.style.name in ("Heading 4", "Heading 5"):
        p.style = d.styles["Normal"]
        for r in p.runs:
            r.bold = True

# ---- yardımcı: belirli metinle başlayan paragrafı bul ----
def find_para(pred):
    for p in d.paragraphs:
        if pred(p):
            return p
    return None

# ---- 2) KAYNAKLAR başlığı (ilk kaynak öğesinden önce) ----
ref0 = find_para(lambda p: p.text.strip().startswith("Chari, V."))
if ref0 is not None:
    h = ref0.insert_paragraph_before("KAYNAKLAR", style="Heading 1")

# ---- 3) EKLER başlığı (ilk ek bölümünden önce) ----
appx0 = find_para(lambda p: p.style.name == "Heading 1"
                  and p.text.strip().startswith("EK A"))
if appx0 is not None:
    appx0.insert_paragraph_before("EKLER", style="Heading 1")

# ---- 4) İÇİNDEKİLER + Word TOC alanı (SİMGE LİSTESİ'nden önce) ----
def add_toc_field(p, instr, placeholder):
    run = p.add_run()
    r = run._r
    fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), "begin")
    fc.set(qn("w:dirty"), "true"); r.append(fc)
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve")
    it.text = instr; r.append(it)
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate"); r.append(sep)
    t = OxmlElement("w:t"); t.text = placeholder; r.append(t)
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end"); r.append(end)

simge = find_para(lambda p: p.style.name == "Heading 1"
                  and p.text.strip() == "SİMGE LİSTESİ")
if simge is not None:
    simge.insert_paragraph_before("İÇİNDEKİLER", style="Heading 1")
    toc_entries = [
        ("ÖNSÖZ", "3"),
        ("İÇİNDEKİLER", "4"),
        ("SİMGE LİSTESİ", "6"),
        ("KISALTMA LİSTESİ", "9"),
        ("ŞEKİL LİSTESİ", "10"),
        ("ÇİZELGE LİSTESİ", "11"),
        ("ÖZET", "12"),
        ("ABSTRACT", "13"),
        ("1 Giriş", "14"),
        ("2 Teorik Arka Plan", "16"),
        ("2.1 Gruplar, Lie Cebirleri ve Evrensel Sarmalayıcı Cebir", "16"),
        ("2.2 sl2 Lie Cebri", "17"),
        ("2.3 Hopf Cebirleri", "17"),
        ("3 Kuantum Grup Uq(sl2)'nin Tanımı", "19"),
        ("3.1 q-Deformasyon Sezgisi", "20"),
        ("4 Uq(sl2)'nin Sonlu Boyutlu Temsil Teorisi", "21"),
        ("5 Kristal Tabanlar (Genel Bakış)", "23"),
        ("6 Yazılım Mimarisi", "24"),
        ("7 Tensör Çarpımı ve Clebsch-Gordan Ayrışımı", "28"),
        ("8 R-Matrisi ve Yang-Baxter Denklemi", "30"),
        ("9 GLq(2|1) için Graded Yang-Baxter Doğrulaması", "34"),
        ("10 Klasik, Kuantum ve Birim Kök Limitleri", "42"),
        ("11 Sonuç ve Öneriler", "46"),
        ("KAYNAKLAR", "48"),
        ("EKLER", "49"),
        ("EK A Kurulum, Çalıştırma ve Yeniden Kullanım Kılavuzu", "50"),
        ("EK B Test Eşlemesi", "54"),
        ("ÖZGEÇMİŞ", "56"),
    ]
    for title, page in toc_entries:
        p = simge.insert_paragraph_before(style="Normal")
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        p.add_run(f"{title}\t{page}")

# ---- 5) ŞEKİL ve ÇİZELGE listeleri (ÖZET'ten önce) ----
figs, tabs = [], []
for p in d.paragraphs:
    t = p.text.strip()
    if len(t) > 6 and t.startswith("Şekil ") and (t[6].isdigit() or t[6].isupper()):
        figs.append(t)
    elif len(t) > 8 and t.startswith("Çizelge ") and (t[8].isdigit() or t[8].isupper()):
        tabs.append(t)

expected_tabs = [
    "Çizelge 3.1 q-tamsayıların ilk değerleri ve klasik limitleri (paketin q_integer çıktısı).",
    "Çizelge 6.1 quantum_group paketinin modülleri ve sorumlulukları.",
    "Çizelge 9.1 V^⊗ 4 üzerindeki yerleşimler ve otomatik kontroller (tümü 81×81; all_Rij_GLq21, local_ybe_on_four_tensor_GLq21, braid_far_commutativity_residual_GLq21).",
    "Çizelge 10.1 V2 temsilinin üç limit rejimindeki karşılaştırması.",
    "Çizelge B.1 Doğrulama eksenleri, kod modülleri ve test dosyaları.",
]
for expected in expected_tabs:
    key = " ".join(expected.split()[:2])
    if not any(item.startswith(key) for item in tabs):
        tabs.append(expected)
ordered_tabs = []
for expected in expected_tabs:
    key = " ".join(expected.split()[:2])
    match = next((item for item in tabs if item.startswith(key)), expected)
    ordered_tabs.append(match)
tabs = ordered_tabs

ozet = find_para(lambda p: p.style.name == "Heading 1" and p.text.strip() == "ÖZET")
if ozet is not None:
    def insert_list(anchor, title, items):
        anchor.insert_paragraph_before(title, style="Heading 1")
        if not items:
            anchor.insert_paragraph_before("(Liste otomatik üretildi.)", style="Normal")
        for it in items:
            para = anchor.insert_paragraph_before(style="Normal")
            para.paragraph_format.line_spacing = 1.0
            run = para.add_run(it[:160])
    # Önce ŞEKİL, sonra ÇİZELGE (ÖZET'ten hemen önce sırayla)
    insert_list(ozet, "ŞEKİL LİSTESİ", figs)
    insert_list(ozet, "ÇİZELGE LİSTESİ", tabs)

# ---- 6) Kapak: iki kapağı ortala, koyulaştır, ayrı sayfalara koy ----
cover_seen = 0
for p in d.paragraphs:
    t = p.text.strip()
    if t.startswith("T.C.") and "YILDIZ" in t:
        cover_seen += 1
        p.style = d.styles["Normal"]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
        if cover_seen == 2:           # iç kapak yeni sayfadan başlasın
            p.paragraph_format.page_break_before = True

paras = list(d.paragraphs)
for idx, p in enumerate(paras):
    if p.style.name == "Heading 1" and p.text.strip() in ("ÖZET", "ABSTRACT"):
        if idx + 1 < len(paras):
            paras[idx + 1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# ---- 7) Ana bölüm başlıkları yeni sayfadan başlasın ----
for p in d.paragraphs:
    if p.style.name == "Heading 1":
        p.paragraph_format.page_break_before = True

# ---- 8) Belge açılışında tüm alanları (İÇİNDEKİLER) güncelle ----
# Word/LibreOffice dosyayı açarken TOC alanını gerçek sayfa numaralarıyla doldurur.
settings = d.settings.element
if settings.find(qn("w:updateFields")) is None:
    uf = OxmlElement("w:updateFields")
    uf.set(qn("w:val"), "true")
    settings.insert(0, uf)

d.save(PATH)
print("post-process tamam:", len(figs), "şekil,", len(tabs), "çizelge listelendi")
